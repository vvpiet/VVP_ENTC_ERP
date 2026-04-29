import streamlit as st
import pandas as pd
import database
from database import *
import os
from datetime import datetime, timedelta
from fpdf import FPDF
import io

COLLEGE_NAME = "Vidya Vikas Pratishthan Institute of Engineering & Technology, Solapur"
DEPARTMENT_NAME = "Department of Electronics and Telecommunication Engineering"
COLLEGE_LOGO_PATH = "college_logo.png"

# Set page config
st.set_page_config(page_title="Lecture Engagement Register", layout="wide")


def render_page_header(subtitle: str = "Lecture Engagement & Resource Management System"):
    with st.container():
        cols = st.columns([1, 8, 1])
        if os.path.exists(COLLEGE_LOGO_PATH):
            cols[0].image(COLLEGE_LOGO_PATH, width=100)
        else:
            cols[0].markdown("<div style='font-size:48px; text-align:center;'>🏫</div>", unsafe_allow_html=True)
        cols[1].markdown(
            f"""
            <div style='text-align:center; padding: 8px 0;'>
                <h2 style='margin: 0; font-size: 28px; color: #0f3c78;'>{COLLEGE_NAME}</h2>
                <p style='margin: 4px 0 4px; font-size: 16px; color: #555;'>{subtitle}</p>
            </div>
            """,
            unsafe_allow_html=True,
        )
        cols[2].write("")
    st.markdown("<hr style='border: 1px solid #ddd; margin: 0 0 16px 0;'>", unsafe_allow_html=True)


def render_page_footer():
    st.markdown("<hr style='border: 1px solid #ddd; margin: 24px 0 8px 0;'>", unsafe_allow_html=True)
    footer_cols = st.columns([1, 8, 1])
    footer_cols[1].markdown(
        """
        <div style='text-align:center; font-size:14px; color:#444;'>
            <strong>Prepared by:</strong> Prof. Amir M. Usman Wagdarikar, Head and Exam Coordinator, Vidya Vikas Pratishthan Institute of Engineering & Technology, Solapur
        </div>
        """,
        unsafe_allow_html=True,
    )

def check_weekly_attendance(student_id, subject_id, weeks=4):
    """Check attendance percentage for last N weeks"""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    start_date = datetime.now() - timedelta(weeks=weeks)
    
    cur.execute(
        "SELECT COUNT(*) as total, SUM(CASE WHEN present THEN 1 ELSE 0 END) as present "
        "FROM attendance "
        "WHERE student_id = %s AND subject_id = %s AND date >= %s",
        (student_id, subject_id, start_date.date())
    )
    result = cur.fetchone()
    cur.close()
    conn.close()
    
    if result['total'] == 0:
        return 0
    
    attendance_percentage = (result['present'] / result['total']) * 100
    return attendance_percentage

def get_defaulter_students(threshold=60, weeks=4):
    """Get students with attendance below threshold"""
    conn = get_db_connection()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    start_date = datetime.now() - timedelta(weeks=weeks)
    
    cur.execute(
        "SELECT DISTINCT s.id, s.name, s.roll_no, c.name as class_name, "
        "sub.id as subject_id, sub.name as subject_name, "
        "(SUM(CASE WHEN a.present THEN 1 ELSE 0 END)::float / COUNT(*)) * 100 as attendance_pct "
        "FROM students s "
        "JOIN classes c ON s.class_id = c.id "
        "JOIN attendance a ON s.id = a.student_id "
        "JOIN subjects sub ON a.subject_id = sub.id "
        "WHERE a.date >= %s "
        "GROUP BY s.id, s.name, s.roll_no, c.name, sub.id, sub.name "
        "HAVING (SUM(CASE WHEN a.present THEN 1 ELSE 0 END)::float / COUNT(*)) * 100 < %s "
        "ORDER BY s.roll_no",
        (start_date.date(), threshold)
    )
    defaulters = cur.fetchall()
    cur.close()
    conn.close()
    return defaulters

def generate_maharashtra_gradecard(student_data, grades_data):
    """Generate Maharashtra University style grade card PDF"""
    
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_margins(12, 10, 12)

    if os.path.exists(COLLEGE_LOGO_PATH):
        pdf.image(COLLEGE_LOGO_PATH, x=12, y=10, w=24)

    # Header
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 8, COLLEGE_NAME, ln=True, align="C")
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 7, student_data.get('department', DEPARTMENT_NAME), ln=True, align="C")
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 5, "Solapur, Maharashtra", ln=True, align="C")
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 7, "STATEMENT OF MARKS / GRADES", ln=True, align="C")

    pdf.ln(2)
    
    # Student Information
    pdf.set_font("Arial", "", 9)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.4)
    
    info_data = [
        [f"Name: {student_data['name']}", f"Roll No: {student_data['roll_no']}"],
        [f"PRN: {student_data.get('prn', '')}", f"Class: {student_data['class']}"],
        [f"Semester: {student_data['semester']}", f"Course: {student_data['course']}"],
        [f"Exam Event: {student_data.get('exam_event', '')}", f"Date: {datetime.now().strftime('%d-%m-%Y')}"]
    ]
    
    for row in info_data:
        pdf.cell(95, 6, row[0], border=1)
        pdf.cell(95, 6, row[1], border=1, ln=True)
    
    pdf.ln(3)
    
    # Table Header
    pdf.set_font("Arial", "B", 8.5)
    col_widths = [9, 42, 11, 11, 11, 10, 10, 10, 10]
    headers = ["Sr", "Subject", "Int", "End", "Total", "Grade", "Cr", "GP", "CP"]
    
    for i, header in enumerate(headers):
        pdf.cell(col_widths[i], 7, header, border=1, align="C")
    pdf.ln()
    
    # Table Data
    pdf.set_font("Arial", "", 8.5)
    total_cp = 0
    total_cr = 0
    
    for i, row in enumerate(grades_data, 1):
        pdf.cell(col_widths[0], 7, str(i), border=1, align="C")
        pdf.cell(col_widths[1], 7, row["subject"][:28], border=1)
        pdf.cell(col_widths[2], 7, str(row["internal"]), border=1, align="C")
        pdf.cell(col_widths[3], 7, str(row["end"]), border=1, align="C")
        pdf.cell(col_widths[4], 7, str(row["total"]), border=1, align="C")
        pdf.cell(col_widths[5], 7, row["grade"], border=1, align="C")
        pdf.cell(col_widths[6], 7, str(row["credits"]), border=1, align="C")
        pdf.cell(col_widths[7], 7, str(row["gp"]), border=1, align="C")
        pdf.cell(col_widths[8], 7, str(row["cp"]), border=1, align="C")
        pdf.ln()
        
        total_cp += row["cp"]
        total_cr += row["credits"]
    
    # Summary
    pdf.ln(4)
    pdf.set_font("Arial", "B", 9)
    sgpa = round(total_cp / total_cr, 2) if total_cr > 0 else 0
    
    pdf.cell(95, 6, f"Total Credits: {total_cr}", border=0)
    pdf.cell(95, 6, f"SGPA: {sgpa}", border=0, ln=True)
    pdf.cell(95, 6, f"Total Credit Points: {total_cp}", border=0)
    pdf.cell(95, 6, f"Result: {'PASS' if sgpa >= 4.0 else 'FAIL'}", border=0, ln=True)
    
    # Grade Scale
    pdf.ln(4)
    pdf.set_font("Arial", "B", 8)
    pdf.cell(0, 5, "GRADE SCALE:", ln=True)
    
    pdf.set_font("Arial", "", 7)
    grades_scale = [
        ["O(10)", "90-100", "A+(9)", "80-89", "A(8)", "70-79"],
        ["B+(7)", "60-69", "B(6)", "50-59", "C(5)", "40-49"],
        ["D(4)", "35-39", "F(0)", "<35", "", ""]
    ]
    
    for grade_row in grades_scale:
        for grade in grade_row:
            pdf.cell(31.5, 4.5, grade, border=1, align="C")
        pdf.ln()
    
    # Footer
    pdf.ln(6)
    pdf.set_font("Arial", "", 8)
    pdf.cell(0, 5, "This is a computer-generated statement and does not require signature.", ln=True, align="C")
    
    pdf.ln(2)
    pdf.set_font("Arial", "", 9)
    pdf.cell(0, 5, "Controller of Examination", align="R", ln=True)
    pdf.cell(0, 5, datetime.now().strftime("%d-%m-%Y"), align="R")
    
    pdf_bytes = pdf.output(dest='S')
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode('latin-1')
    elif isinstance(pdf_bytes, bytearray):
        pdf_bytes = bytes(pdf_bytes)
    return pdf_bytes


def generate_gradecard_docx(student_data, grades_data):
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        raise ImportError("python-docx is required to generate Word documents. Install it with 'pip install python-docx'.")

    doc = Document()
    doc.add_heading('Grade Card', level=1)
    doc.add_paragraph(f"Name: {student_data['name']}")
    doc.add_paragraph(f"Roll No: {student_data['roll_no']}")
    doc.add_paragraph(f"PRN: {student_data.get('prn', '')}")
    doc.add_paragraph(f"Class: {student_data['class']}")
    doc.add_paragraph(f"Semester: {student_data['semester']}")
    doc.add_paragraph(f"Exam Event: {student_data.get('exam_event', '')}")
    doc.add_paragraph(f"Course: {student_data['course']}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")

    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Sr', 'Subject', 'Int', 'End', 'Total', 'Grade', 'Cr', 'GP', 'CP']
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header

    total_cp = 0
    total_cr = 0
    for idx, row in enumerate(grades_data, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = row['subject']
        row_cells[2].text = str(row['internal'])
        row_cells[3].text = str(row['end'])
        row_cells[4].text = str(row['total'])
        row_cells[5].text = row['grade']
        row_cells[6].text = str(row['credits'])
        row_cells[7].text = str(row['gp'])
        row_cells[8].text = str(row['cp'])
        total_cp += row['cp']
        total_cr += row['credits']

    doc.add_paragraph('')
    doc.add_paragraph(f'Total Credits: {total_cr}')
    sgpa = round(total_cp / total_cr, 2) if total_cr > 0 else 0
    doc.add_paragraph(f'SGPA: {sgpa}')
    doc.add_paragraph(f'Total Credit Points: {total_cp}')
    doc.add_paragraph(f"Result: {'PASS' if sgpa >= 4.0 else 'FAIL'}")

    doc.add_paragraph('')
    doc.add_paragraph('Grade Scale:')
    doc.add_paragraph('O(10): 90-100   A+(9): 80-89   A(8): 70-79')
    doc.add_paragraph('B+(7): 60-69   B(6): 50-59   C(5): 40-49')
    doc.add_paragraph('D(4): 35-39   F(0): <35')

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# Initialize database
if 'db_init' not in st.session_state:
    create_tables()
    ensure_schema()
    # Insert initial classes
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("INSERT INTO classes (name) VALUES ('SY'), ('TY'), ('B.Tech') ON CONFLICT (name) DO NOTHING")
    # Create default admin
    try:
        create_user('admin', 'admin123', 'admin', 'Administrator', 'admin@example.com')
    except:
        pass
    conn.commit()
    cur.close()
    conn.close()
    st.session_state.db_init = True
else:
    ensure_schema()

# Login function
def login():
    render_page_header()
    st.title("Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login", key="login_button"):
        user = authenticate_user(username, password)
        if user:
            st.session_state.user = user
            st.success("Logged in successfully")
            st.rerun()
        else:
            st.error("Invalid credentials")
    render_page_footer()

# Logout
def logout():
    if 'user' in st.session_state:
        del st.session_state.user
    st.rerun()

# Admin page
def admin_page():
    render_page_header("Admin Portal: Manage subjects, faculty assignments, attendance, and engagement reports")
    st.title("Admin Dashboard")
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10 = st.tabs(["Upload Students", "Manage Subjects", "Assign Faculty", "Download Attendance", "Download Engagement", "MCQ Reports", "Create Users", "Manage Students", "Defaulter Students", "Generate Grade Card"])
    
    with tab1:
        st.header("Upload Student List")
        uploaded_file = st.file_uploader("Upload CSV (roll_no, name, class_name, prn)", type="csv")
        if uploaded_file:
            df = pd.read_csv(uploaded_file)
            if st.button("Upload", key="upload_students"):
                conn = get_db_connection()
                cur = conn.cursor()
                for _, row in df.iterrows():
                    # Get class_id
                    cur.execute("SELECT id FROM classes WHERE name = %s", (row['class_name'],))
                    class_id = cur.fetchone()[0]
                    prn_value = row['prn'] if 'prn' in row and not pd.isna(row['prn']) else None
                    cur.execute("INSERT INTO students (roll_no, prn, name, class_id) VALUES (%s, %s, %s, %s) ON CONFLICT (roll_no) DO NOTHING",
                                (row['roll_no'], prn_value, row['name'], class_id))
                    # Create user
                    cur.execute("INSERT INTO users (username, password_hash, role, name, email) VALUES (%s, %s, 'student', %s, '') ON CONFLICT (username) DO NOTHING",
                                (row['roll_no'], hash_password('student123'), row['name']))
                conn.commit()
                cur.close()
                conn.close()
                st.success("Students uploaded")
    
    with tab2:
        st.header("Add Subjects to Classes")
        class_name = st.selectbox("Class", ["SY", "TY", "B.Tech"], key="admin_add_subject_class")
        subject_name = st.text_input("Subject Name")
        if st.button("Add Subject", key="add_subject"):
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute("SELECT id FROM classes WHERE name = %s", (class_name,))
            class_id = cur.fetchone()[0]
            cur.execute("INSERT INTO subjects (name, class_id) VALUES (%s, %s)", (subject_name, class_id))
            conn.commit()
            cur.close()
            conn.close()
            st.success("Subject added")
    
    with tab3:
        st.header("Assign Subjects to Faculty")
        # List faculty
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT id, name FROM users WHERE role = 'faculty'")
        faculty = cur.fetchall()
        faculty_names = [f['name'] for f in faculty]
        if not faculty_names:
            st.warning("No faculty available to assign subjects")
        else:
            faculty_dict = {f['name']: f['id'] for f in faculty}
            selected_faculty = st.selectbox("Faculty", faculty_names, key="admin_assign_faculty")
            faculty_id = faculty_dict[selected_faculty]
            
            # List subjects
            cur.execute("SELECT s.id, s.name, c.name as class_name FROM subjects s JOIN classes c ON s.class_id = c.id")
            subjects = cur.fetchall()
            subject_options = [f"{s['name']} ({s['class_name']})" for s in subjects]
            if not subject_options:
                st.warning("No subjects available")
            else:
                subject_dict = {f"{s['name']} ({s['class_name']})": s['id'] for s in subjects}
                selected_subject = st.selectbox("Subject", subject_options, key="admin_assign_subject")
                subject_id = subject_dict[selected_subject]
                
                if st.button("Assign", key="assign_subject"):
                    cur.execute("INSERT INTO faculty_subjects (faculty_id, subject_id) VALUES (%s, %s) ON CONFLICT DO NOTHING", (faculty_id, subject_id))
                    conn.commit()
                    st.success("Assigned")
        cur.close()
        conn.close()
    
    with tab4:
        st.header("Download Monthly Attendance")
        month = st.selectbox("Month", list(range(1,13)), key="admin_attendance_month")
        year = st.number_input("Year", value=2023, key="admin_attendance_year")
        if st.button("Download", key="download_attendance"):
            # Query attendance
            conn = get_db_connection()
            df = pd.read_sql(f"SELECT s.roll_no, s.name, c.name as class, sub.name as subject, a.date, a.time, a.present FROM attendance a JOIN students s ON a.student_id = s.id JOIN subjects sub ON a.subject_id = sub.id JOIN classes c ON s.class_id = c.id WHERE EXTRACT(MONTH FROM a.date) = {month} AND EXTRACT(YEAR FROM a.date) = {year}", conn)
            conn.close()
            csv = df.to_csv(index=False)
            st.download_button("Download CSV", csv, "attendance.csv", key="download_att_csv")
    
    with tab5:
        st.header("Download Lecture Engagement")
        period = st.selectbox("Period", ["Weekly", "Monthly"], key="admin_engagement_period")
        if period == "Weekly":
            week_start = st.date_input("Week Start", key="admin_engagement_week_start")
            # Calculate week end
            week_end = week_start + pd.Timedelta(days=6)
        else:
            month = st.selectbox("Month", list(range(1,13)), key="admin_engagement_month")
            year = st.number_input("Year", value=2023, key="admin_engagement_year")
        if st.button("Download", key="download_engagement"):
            conn = get_db_connection()
            if period == "Weekly":
                cur = conn.cursor(cursor_factory=RealDictCursor)
                cur.execute(
                    "SELECT le.date, u.name as faculty, s.name as subject, le.topic_covered, le.lecture_number, le.syllabus_percent, le.total_present, le.total_absent, array_to_string(le.absent_roll_numbers, ', ') as absent_roll_numbers FROM lecture_engagement le JOIN users u ON le.faculty_id = u.id JOIN subjects s ON le.subject_id = s.id WHERE le.date BETWEEN %s AND %s",
                    (week_start, week_end)
                )
                rows = cur.fetchall()
                cur.close()
                df = pd.DataFrame(rows)
            else:
                df = pd.read_sql(f"SELECT le.date, u.name as faculty, s.name as subject, le.topic_covered, le.lecture_number, le.syllabus_percent, le.total_present, le.total_absent, array_to_string(le.absent_roll_numbers, ', ') as absent_roll_numbers FROM lecture_engagement le JOIN users u ON le.faculty_id = u.id JOIN subjects s ON le.subject_id = s.id WHERE EXTRACT(MONTH FROM le.date) = {month} AND EXTRACT(YEAR FROM le.date) = {year}", conn)
            conn.close()
            csv = df.to_csv(index=False)
            st.download_button("Download CSV", csv, "engagement.csv", key="download_eng_csv")
    
    with tab6:
        st.header("MCQ Reports")
        st.write("View all MCQ test results submitted by students.")
        report = get_mcq_test_results()
        if not report:
            st.info("No MCQ test results available yet.")
        else:
            df_report = pd.DataFrame(report)
            st.dataframe(df_report, use_container_width=True)
            csv = df_report.to_csv(index=False)
            st.download_button("Download MCQ Results CSV", csv, "mcq_test_results.csv", key="download_mcq_results")
    
    with tab7:
        st.header("Create Users")
        role = st.selectbox("Role", ["faculty", "student"], key="admin_create_user_role")
        username = st.text_input("Username (login ID)", key="admin_create_username")
        password = st.text_input("Password", type="password", key="admin_create_password")
        name = st.text_input("Name", key="admin_create_name")
        email = st.text_input("Email", key="admin_create_email")
        roll_no = ""
        prn = ""
        if role == "student":
            roll_no = st.text_input("Student Roll Number", key="admin_create_user_roll_no")
            st.info("Student login username will default to roll number if left blank.")
            class_name = st.selectbox("Class", ["SY", "TY", "B.Tech"], key="admin_create_user_class")
            prn = st.text_input("PRN", key="admin_create_user_prn")
        if st.button("Create", key="create_user"):
            try:
                if role == "student":
                    if not roll_no:
                        st.error("Student Roll Number is required.")
                    else:
                        username_to_use = username.strip() or roll_no.strip()
                        conn = get_db_connection()
                        cur = conn.cursor()
                        cur.execute("SELECT id FROM classes WHERE name = %s", (class_name,))
                        class_id = cur.fetchone()[0]
                        cur.execute("INSERT INTO students (roll_no, prn, name, class_id) VALUES (%s, %s, %s, %s) ON CONFLICT (roll_no) DO NOTHING",
                                    (roll_no, prn or None, name, class_id))
                        conn.commit()
                        cur.close()
                        conn.close()
                        create_user(username_to_use, password, role, name, email)
                        st.success(f"Student user created with roll number {roll_no}.")
                else:
                    if not username:
                        st.error("Username is required for faculty users.")
                    else:
                        create_user(username, password, role, name, email)
                        st.success("Faculty user created")
            except Exception as e:
                st.error(f"Error creating user: {str(e)}")
    
    with tab8:
        st.header("Manage Students")
        students = get_all_students()
        if not students:
            st.warning("No students found")
        else:
            student_options = [f"{s['roll_no']} - {s['name']} ({s['class_name']})" for s in students]
            student_map = {option: s['id'] for option, s in zip(student_options, students)}
            selected_student_options = st.multiselect("Select students to delete", student_options, key="selected_students_to_delete")
            if st.button("Delete selected students", key="delete_selected_students"):
                if selected_student_options:
                    delete_students([student_map[opt] for opt in selected_student_options])
                    st.success(f"Deleted {len(selected_student_options)} students.")
                    st.experimental_rerun()
                else:
                    st.warning("Select one or more students before deleting.")
            st.divider()
            st.subheader("Student List")
            for student in students:
                col1, col2, col3, col4, col5, col6 = st.columns([2, 2, 2, 2, 1, 1])
                with col1:
                    st.write(student['roll_no'])
                with col2:
                    st.write(student['name'])
                with col3:
                    st.write(student['class_name'])
                with col4:
                    st.write(student.get('prn', ''))
                with col5:
                    if st.button("Edit", key=f"edit_student_{student['id']}"):
                        st.session_state.edit_student_id = student['id']
                with col6:
                    if st.button("Delete", key=f"delete_student_{student['id']}"):
                        st.session_state.delete_student_id = student['id']
            
            if 'edit_student_id' in st.session_state:
                st.divider()
                st.subheader("Edit Student")
                student = next((s for s in students if s['id'] == st.session_state.edit_student_id), None)
                if student:
                    new_roll_no = st.text_input("Roll No", value=student['roll_no'], key="edit_roll_no")
                    new_name = st.text_input("Name", value=student['name'], key="edit_name")
                    new_prn = st.text_input("PRN", value=student.get('prn', ''), key="edit_prn")
                    new_class = st.selectbox("Class", ["SY", "TY", "B.Tech"], index=["SY", "TY", "B.Tech"].index(student['class_name']), key="edit_class")
                    if st.button("Save Changes", key="save_student_changes"):
                        update_student(st.session_state.edit_student_id, new_roll_no, new_prn, new_name, new_class)
                        del st.session_state.edit_student_id
                        st.success("Student updated")
                        st.rerun()
            
            if 'delete_student_id' in st.session_state:
                st.divider()
                st.warning("Are you sure you want to delete this student? All related attendance records will be deleted.")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Confirm Delete", key="confirm_delete_student"):
                        delete_student(st.session_state.delete_student_id)
                        del st.session_state.delete_student_id
                        st.success("Student deleted")
                        st.rerun()
                with col2:
                    if st.button("Cancel", key="cancel_delete_student"):
                        del st.session_state.delete_student_id
                        st.rerun()
    
    with tab9:
        st.header("Defaulter Students (Attendance < 60%)")
        weeks = st.slider("Check last N weeks", 1, 16, 4, key="defaulter_weeks")
        
        if st.button("Get Defaulter List", key="get_defaulters"):
            defaulters = get_defaulter_students(threshold=60, weeks=weeks)
            
            if not defaulters:
                st.success(f"✅ No defaulter students found in last {weeks} weeks!")
            else:
                st.warning(f"⚠️ Found {len(defaulters)} defaulter student(s)")
                
                # Display in table format
                df_defaulters = pd.DataFrame([
                    {
                        'Roll No': d['roll_no'],
                        'Name': d['name'],
                        'Class': d['class_name'],
                        'Subject': d['subject_name'],
                        'Attendance %': f"{d['attendance_pct']:.2f}%"
                    }
                    for d in defaulters
                ])
                st.dataframe(df_defaulters, use_container_width=True)
                
                # Send alerts
                if st.button("Send Alert Messages to Defaulters", key="send_alerts"):
                    try:
                        for defaulter in defaulters:
                            # You can integrate with WhatsApp API (Twilio) here
                            st.info(f"Alert would be sent to {defaulter['name']} ({defaulter['roll_no']}) - Attendance: {defaulter['attendance_pct']:.2f}%")
                        st.success("Alert notifications would be sent to all defaulter students!")
                    except Exception as e:
                        st.error(f"Error sending alerts: {str(e)}")

    with tab10:
        st.header("Generate Grade Card")
        students = get_all_students()
        if not students:
            st.warning("No students available")
        else:
            student_options = [f"{s['roll_no']} - {s['name']} ({s['class_name']})" for s in students]
            student_dict = {option: s for option, s in zip(student_options, students)}
            selected_student = st.selectbox("Select Student", student_options, key="admin_gradecard_student")
            student = student_dict[selected_student]

            st.write(f"**Name:** {student['name']}")
            st.write(f"**Roll No:** {student['roll_no']}")
            st.write(f"**Class:** {student['class_name']}")

            semester = st.text_input("Semester", value="IV", key="admin_gradecard_semester")
            course = st.text_input("Course", value="B.Tech", key="admin_gradecard_course")
            exam_term = st.selectbox("Exam Term", ["Summer", "Winter"], key="admin_gradecard_term")
            exam_type = st.selectbox("Exam Type", ["Regular", "Supplementary"], key="admin_gradecard_type")
            exam_year = st.text_input("Exam Year", value=str(datetime.now().year), key="admin_gradecard_year")
            subject_count = st.number_input("Number of subjects", 1, 10, 6, key="admin_gradecard_subject_count")
            grades_data = []

            for i in range(1, subject_count + 1):
                col1, col2, col3, col4, col5 = st.columns(5)
                with col1:
                    subject = st.text_input(f"Subject Name {i}", key=f"admin_subject_{i}")
                with col2:
                    subject_type = st.selectbox(f"Subject Type {i}", ["Theory", "Practical"], key=f"admin_subject_type_{i}")
                with col3:
                    if subject_type == "Practical":
                        internal = st.number_input(f"Internal {i} (out of 60)", 0, 60, key=f"admin_internal_{i}")
                    else:
                        internal = st.number_input(f"Internal {i} (out of 40)", 0, 40, key=f"admin_internal_{i}")
                with col4:
                    if subject_type == "Practical":
                        end = st.number_input(f"External {i} (out of 40)", 0, 40, key=f"admin_end_{i}")
                    else:
                        end = st.number_input(f"External {i} (out of 60)", 0, 60, key=f"admin_end_{i}")
                with col5:
                    credits = st.number_input(f"Credits {i}", 1, 6, 4, key=f"admin_credits_{i}")
                    total = internal + end
                    if total >= 91:
                        grade = "EX"
                        gp = 10
                    elif total >= 86:
                        grade = "AA"
                        gp = 9
                    elif total >= 81:
                        grade = "AB"
                        gp = 8.5
                    elif total >= 76:
                        grade = "BB"
                        gp = 8
                    elif total >= 71:
                        grade = "BC"
                        gp = 7.5
                    elif total >= 66:
                        grade = "CC"
                        gp = 7
                    elif total >= 61:
                        grade = "CD"
                        gp = 6.5
                    elif total >= 56:
                        grade = "DD"
                        gp = 6
                    elif total >= 51:
                        grade = "DE"
                        gp = 5.5
                    elif total >= 40:
                        grade = "EE"
                        gp = 5
                    else:
                        grade = "F"
                        gp = 0
                    st.write(f"**Grade: {grade}**")

                if subject:
                    grades_data.append({
                        "subject": subject,
                        "internal": internal,
                        "end": end,
                        "total": total,
                        "grade": grade,
                        "credits": credits,
                        "gp": gp,
                        "cp": gp * credits
                    })

            if grades_data and st.button("Generate and Save Grade Card", key="admin_generate_gradecard"):
                student_data = {
                    "name": student['name'],
                    "prn": student.get('prn', ''),
                    "roll_no": student['roll_no'],
                    "class": student['class_name'],
                    "semester": semester,
                    "course": course,
                    "exam_event": f"{exam_term} {exam_year} {exam_type}",
                    "department": DEPARTMENT_NAME
                }
                pdf_bytes = generate_maharashtra_gradecard(student_data, grades_data)
                save_gradecard(student['id'], pdf_bytes, semester, course)
                st.success("✅ Grade card generated and saved for student download.")
                st.download_button(
                    "📥 Download Generated Grade Card",
                    pdf_bytes,
                    f"gradecard_{student['roll_no']}.pdf",
                    "application/pdf",
                    key="admin_download_gradecard"
                )
                try:
                    word_bytes = generate_gradecard_docx(student_data, grades_data)
                    st.download_button(
                        "📥 Download Generated Grade Card DOCX",
                        word_bytes,
                        f"gradecard_{student['roll_no']}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="admin_download_gradecard_docx"
                    )
                except ImportError as e:
                    st.warning(str(e))

    render_page_footer()

# Faculty page
def faculty_page():
    render_page_header("Faculty Portal: Mark attendance, submit lecture engagement, and upload resources")
    st.title("Faculty Dashboard")
    user = st.session_state.user
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Mark Attendance", "Lecture Engagement", "MCQ Tests", "Upload Resources", "View Resources"])
    
    with tab1:
        st.header("Mark Attendance")
        # Get assigned subjects
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT s.id, s.name, c.name as class_name FROM subjects s JOIN faculty_subjects fs ON s.id = fs.subject_id JOIN classes c ON s.class_id = c.id WHERE fs.faculty_id = %s", (user['id'],))
        subjects = cur.fetchall()
        subject_options = [f"{s['name']} ({s['class_name']})" for s in subjects]
        if not subject_options:
            st.warning("No subjects assigned to you")
        else:
            subject_dict = {f"{s['name']} ({s['class_name']})": s['id'] for s in subjects}
            selected_subject = st.selectbox("Subject", subject_options, key="faculty_att_subject")
            subject_id = subject_dict[selected_subject]
            
            date = st.date_input("Date", key="att_date")
            time = st.time_input("Time", key="att_time")
            
            # Get students for the class
            cur.execute("SELECT st.id, st.roll_no, st.name FROM students st JOIN subjects sub ON st.class_id = sub.class_id WHERE sub.id = %s", (subject_id,))
            students = cur.fetchall()
            cur.close()
            conn.close()
            
            attendance = {}
            for student in students:
                attendance[student['id']] = st.checkbox(f"{student['roll_no']} - {student['name']}", key=f"att_{student['id']}")
            
            if st.button("Submit Attendance", key="submit_attendance"):
                conn = get_db_connection()
                cur = conn.cursor()
                for student_id, present in attendance.items():
                    cur.execute("INSERT INTO attendance (student_id, subject_id, faculty_id, date, time, present) VALUES (%s, %s, %s, %s, %s, %s)",
                                (student_id, subject_id, user['id'], date, time, present))
                conn.commit()
                cur.close()
                conn.close()
                st.success("Attendance marked")
    
    with tab2:
        st.header("Lecture Engagement Register")
        # Similar to above, select subject, date
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT s.id, s.name, c.name as class_name FROM subjects s JOIN faculty_subjects fs ON s.id = fs.subject_id JOIN classes c ON s.class_id = c.id WHERE fs.faculty_id = %s", (user['id'],))
        subjects = cur.fetchall()
        subject_options = [f"{s['name']} ({s['class_name']})" for s in subjects]
        if not subject_options:
            st.warning("No subjects assigned to you")
        else:
            subject_dict = {f"{s['name']} ({s['class_name']})": s['id'] for s in subjects}
            selected_subject = st.selectbox("Subject", subject_options, key="eng_subject")
            subject_id = subject_dict[selected_subject]
            
            date = st.date_input("Date", key="eng_date")
            eng_time = st.time_input("Lecture Time", key="eng_time")
            topic = st.text_area("Topic Covered")
            lecture_num = st.number_input("Lecture Number", min_value=1)
            syllabus_pct = st.number_input("% Syllabus Covered", min_value=0.0, max_value=100.0)
            
            # Get attendance records for that SPECIFIC lecture (date + time)
            cur.execute("SELECT st.roll_no, a.present FROM attendance a JOIN students st ON a.student_id = st.id WHERE a.subject_id = %s AND a.faculty_id = %s AND a.date = %s AND a.time = %s ORDER BY st.roll_no", (subject_id, user['id'], date, eng_time))
            att_records = cur.fetchall()
            total_students = len(att_records)
            present = sum(1 for a in att_records if a['present'])
            absent = total_students - present
            absent_roll_numbers = [a['roll_no'] for a in att_records if not a['present']]
            
            st.write(f"Total Students: {total_students}, Present: {present}, Absent: {absent}")
            st.info(f"Absent Students: {', '.join(absent_roll_numbers) if absent_roll_numbers else 'None'}")
            
            if st.button("Submit Engagement", key="submit_engagement"):
                database.store_lecture_engagement(user['id'], subject_id, date, topic, lecture_num, syllabus_pct, present, absent, absent_roll_numbers)
                st.success("Submitted")
        cur.close()
        conn.close()
    
    with tab3:
        st.header("MCQ Test Management")
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT s.id, s.name, c.name as class_name FROM subjects s JOIN faculty_subjects fs ON s.id = fs.subject_id JOIN classes c ON s.class_id = c.id WHERE fs.faculty_id = %s", (user['id'],))
        subjects = cur.fetchall()
        cur.close()
        conn.close()
        if not subjects:
            st.warning("No subjects assigned to you")
        else:
            subject_dict = {f"{s['name']} ({s['class_name']})": s['id'] for s in subjects}
            subject_options = list(subject_dict.keys())
            selected_subject = st.selectbox("Subject", subject_options, key="mcq_subject")
            subject_id = subject_dict[selected_subject]
            test_title = st.text_input("Test Title", key="mcq_test_title")
            proctor_notes = st.text_area("Proctor Notes / Proctoring Instructions", key="mcq_proctor_notes")
            num_questions = st.number_input("Number of Questions", 1, 20, key="mcq_num_questions")

            questions = []
            for q_index in range(1, num_questions + 1):
                st.markdown(f"**Question {q_index}**")
                question_text = st.text_area(f"Question {q_index}", key=f"mcq_question_{q_index}")
                option_a = st.text_input(f"Option A", key=f"mcq_option_a_{q_index}")
                option_b = st.text_input(f"Option B", key=f"mcq_option_b_{q_index}")
                option_c = st.text_input(f"Option C", key=f"mcq_option_c_{q_index}")
                option_d = st.text_input(f"Option D", key=f"mcq_option_d_{q_index}")
                correct_option = st.selectbox("Correct Option", ["A", "B", "C", "D"], key=f"mcq_correct_{q_index}")
                marks = st.number_input("Marks", 1, 20, 1, key=f"mcq_marks_{q_index}")
                questions.append({
                    'question_text': question_text.strip(),
                    'option_a': option_a.strip(),
                    'option_b': option_b.strip(),
                    'option_c': option_c.strip(),
                    'option_d': option_d.strip(),
                    'correct_option': correct_option,
                    'marks': marks
                })

            if st.button("Save MCQ Test", key="save_mcq_test"):
                missing = [i + 1 for i, q in enumerate(questions) if not q['question_text'] or not q['option_a'] or not q['option_b'] or not q['option_c'] or not q['option_d']]
                if not test_title.strip():
                    st.error("Test title is required.")
                elif missing:
                    st.error(f"Please complete all fields for question(s): {', '.join(map(str, missing))}.")
                else:
                    test_id = create_mcq_test(user['id'], subject_id, test_title.strip(), proctor_notes.strip())
                    for q in questions:
                        add_mcq_question(test_id, q['question_text'], q['option_a'], q['option_b'], q['option_c'], q['option_d'], q['correct_option'], q['marks'])
                    st.success("MCQ test created successfully.")

        st.divider()
        st.subheader("Your MCQ Tests")
        tests = get_faculty_tests(user['id'])
        if not tests:
            st.info("No MCQ tests created yet.")
        else:
            for test in tests:
                st.write(f"**{test['title']}** ({test['subject_name']}) - Created: {test['created_at'].strftime('%Y-%m-%d')}")
                st.write(f"Proctor Notes: {test['proctor_notes'] or 'None'}")
                st.markdown("---")
    
    with tab4:
        st.header("Upload Resources (Assignments & Notes)")
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT s.id, s.name, c.name as class_name FROM subjects s JOIN faculty_subjects fs ON s.id = fs.subject_id JOIN classes c ON s.class_id = c.id WHERE fs.faculty_id = %s", (user['id'],))
        subjects = cur.fetchall()
        cur.close()
        conn.close()
        
        subject_options = [f"{s['name']} ({s['class_name']})" for s in subjects]
        if not subject_options:
            st.warning("No subjects assigned to you")
        else:
            subject_dict = {f"{s['name']} ({s['class_name']})": s['id'] for s in subjects}
            selected_subject = st.selectbox("Subject", subject_options, key="upload_subject")
            subject_id = subject_dict[selected_subject]
            
            resource_type = st.selectbox("Resource Type", ["Assignment", "Notes", "Other"], key="resource_type")
            uploaded_file = st.file_uploader("Choose file", key="resource_file")
            
            if uploaded_file and st.button("Upload", key="upload_resource"):
                file_data = uploaded_file.read()
                database.upload_resource(user['id'], subject_id, uploaded_file.name, file_data, resource_type)
                st.success(f"{resource_type} uploaded successfully")
                st.rerun()
    
    with tab5:
        st.header("Your Resources")
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get resources uploaded by this faculty
        cur.execute(
            "SELECT r.id, r.file_name, r.resource_type, s.name as subject_name, r.uploaded_date "
            "FROM resources r "
            "JOIN subjects s ON r.subject_id = s.id "
            "WHERE r.faculty_id = %s "
            "ORDER BY r.uploaded_date DESC",
            (user['id'],)
        )
        resources = cur.fetchall()
        cur.close()
        conn.close()
        
        if not resources:
            st.info("No resources uploaded yet")
        else:
            for resource in resources:
                col1, col2, col3, col4, col5 = st.columns([2, 2, 2, 1, 1])
                with col1:
                    st.write(resource['file_name'])
                with col2:
                    st.write(resource['resource_type'])
                with col3:
                    st.write(resource['subject_name'])
                with col4:
                    st.write(resource['uploaded_date'])
                with col5:
                    if st.button("Delete", key=f"delete_resource_{resource['id']}"):
                        conn = get_db_connection()
                        cur = conn.cursor()
                        cur.execute("DELETE FROM resources WHERE id = %s AND faculty_id = %s", (resource['id'], user['id']))
                        conn.commit()
                        cur.close()
                        conn.close()
                        st.success("Resource deleted")
                        st.rerun()
    render_page_footer()

# Student page
def student_page():
    render_page_header("Student Portal: View attendance, download faculty resources, and download grade card")
    st.title("Student Dashboard")
    user = st.session_state.user
    tab1, tab2, tab3, tab4 = st.tabs(["View Attendance", "Download Resources", "MCQ Tests", "Download Grade Card"])
    
    with tab1:
        st.header("Your Attendance")
        conn = get_db_connection()
        df = pd.read_sql("SELECT sub.name as subject, a.date, a.time, a.present FROM attendance a JOIN subjects sub ON a.subject_id = sub.id WHERE a.student_id = (SELECT id FROM students WHERE roll_no = %s) ORDER BY a.date DESC", conn, params=(user['username'],))
        conn.close()
        
        if df.empty:
            st.info("No attendance records found")
        else:
            # Calculate attendance percentage by subject
            st.subheader("Attendance Summary")
            conn = get_db_connection()
            cur = conn.cursor(cursor_factory=RealDictCursor)
            cur.execute(
                "SELECT sub.name as subject, "
                "COUNT(*) as total, "
                "SUM(CASE WHEN a.present THEN 1 ELSE 0 END) as present, "
                "ROUND(100.0 * SUM(CASE WHEN a.present THEN 1 ELSE 0 END)::numeric / NULLIF(COUNT(*), 0), 2) as attendance_pct "
                "FROM attendance a "
                "JOIN subjects sub ON a.subject_id = sub.id "
                "WHERE a.student_id = (SELECT id FROM students WHERE roll_no = %s) "
                "GROUP BY sub.name "
                "ORDER BY sub.name",
                (user['username'],)
            )
            summary = cur.fetchall()
            cur.close()
            conn.close()
            
            for row in summary:
                color = "🟢" if row['attendance_pct'] >= 75 else "🟡" if row['attendance_pct'] >= 60 else "🔴"
                st.write(f"{color} {row['subject']}: {row['attendance_pct']}% ({row['present']}/{row['total']})")
            
            st.subheader("Detailed Attendance")
            st.dataframe(df, use_container_width=True)

    with tab2:
        st.header("Available Resources")
        try:
            resources = database.get_student_resources(user['username'])
        except Exception as e:
            st.error(f"Error loading resources: {str(e)}")
            resources = []

        if not resources:
            st.info("No resources available yet")
        else:
            for resource in resources:
                cols = st.columns([2, 2, 2, 2, 1])
                with cols[0]:
                    st.write(resource['file_name'])
                with cols[1]:
                    st.write(resource['subject_name'])
                with cols[2]:
                    st.write(resource['resource_type'])
                with cols[3]:
                    st.write(resource['uploaded_date'])
                with cols[4]:
                    st.download_button("Download", data=bytes(resource['file_data']), file_name=resource['file_name'], key=f"download_resource_{resource['id']}")

    with tab3:
        st.header("MCQ Tests")
        tests = get_student_tests(user['username'])
        if not tests:
            st.info("No MCQ tests available yet.")
        else:
            test_options = [f"{t['title']} ({t['subject_name']})" for t in tests]
            selected_test = st.selectbox("Select Test", test_options, key="student_mcq_test")
            if selected_test:
                selected_index = test_options.index(selected_test)
                test_info = tests[selected_index]
                test_data = get_test_with_questions(test_info['id'])
                st.write(f"**{test_data['test']['title']}**")
                st.write(f"Subject: {test_data['test']['subject_name']}")
                st.write(f"Faculty: {test_data['test']['faculty_name']}")
                st.write(f"Proctor Notes: {test_data['test']['proctor_notes'] or 'None'}")
                st.write(f"Proctored: {'Yes' if test_data['test']['proctored'] else 'No'}")

                answers = {}
                for q in test_data['questions']:
                    st.markdown(f"**Q{q['id']}: {q['question_text']}**")
                    choice = st.radio(
                        "Select answer",
                        [f"A. {q['option_a']}", f"B. {q['option_b']}", f"C. {q['option_c']}", f"D. {q['option_d']}"],
                        key=f"student_mcq_q_{q['id']}"
                    )
                    answers[q['id']] = choice[0]

                if st.button("Submit Test", key="submit_student_mcq"):
                    total_marks = 0
                    score = 0
                    answer_records = []
                    for q in test_data['questions']:
                        selected_option = answers.get(q['id'])
                        correct = (selected_option == q['correct_option'])
                        total_marks += q['marks']
                        if correct:
                            score += q['marks']
                        answer_records.append({
                            'question_id': q['id'],
                            'selected_option': selected_option,
                            'is_correct': correct
                        })
                    percent = round((score / total_marks) * 100, 2) if total_marks else 0
                    passed = percent >= 50
                    submit_student_test_attempt(test_info['id'], user['username'], answer_records, score, total_marks, percent, passed, test_data['test']['proctor_notes'])
                    st.success(f"Test submitted. Score: {score}/{total_marks} ({percent}%). Result: {'PASS' if passed else 'FAIL'}")

            attempts = get_student_test_attempts(user['username'])
            if attempts:
                st.divider()
                st.subheader("Your Test Results")
                for attempt in attempts:
                    st.write(f"**{attempt['title']}** - {attempt['subject_name']} | Score: {attempt['score']}/{attempt['total_marks']} | {attempt['percent']}% | {'PASS' if attempt['passed'] else 'FAIL'} | {attempt['finished_at'].strftime('%Y-%m-%d %H:%M')}")

    with tab4:
        st.header("Download Grade Card")
        
        # Get student details
        conn = get_db_connection()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            "SELECT s.id, s.prn, s.name, s.roll_no, c.name as class_name, u.role "
            "FROM students s "
            "JOIN classes c ON s.class_id = c.id "
            "JOIN users u ON u.username = %s "
            "WHERE s.roll_no = %s",
            (user['username'], user['username'])
        )
        student_info = cur.fetchone()
        cur.close()
        conn.close()
        
        if student_info:
            st.write(f"**Name:** {student_info['name']}")
            st.write(f"**Roll No:** {student_info['roll_no']}")
            st.write(f"**PRN:** {student_info.get('prn', '')}")
            st.write(f"**Class:** {student_info['class_name']}")
            
            gradecard_record = get_gradecard(student_info['id'])
            if gradecard_record:
                st.success("✅ Admin-generated grade card is available for download.")
                st.download_button(
                    "📥 Download Approved Grade Card",
                    gradecard_record['pdf_file'],
                    f"gradecard_{student_info['roll_no']}.pdf",
                    "application/pdf",
                    key="download_stored_gradecard"
                )
                st.info("If you want to regenerate the grade card manually, fill the form below.")
            else:
                st.info("No admin-generated grade card found yet. You can create one manually below.")

            st.subheader("Enter Your Grades")
            
            num_subjects = st.number_input("Number of Subjects", 1, 10, 4, key="num_subjects")
            
            grades_data = []
            for i in range(num_subjects):
                st.markdown(f"**Subject {i+1}**")
                col1, col2, col3, col4, col5 = st.columns(5)
                
                with col1:
                    subject = st.text_input(f"Subject Name {i}", key=f"subject_{i}")
                with col2:
                    subject_type = st.selectbox(f"Subject Type {i}", ["Theory", "Practical"], key=f"subject_type_{i}")
                with col3:
                    if subject_type == "Practical":
                        internal = st.number_input(f"Internal {i} (out of 60)", 0, 60, key=f"internal_{i}")
                    else:
                        internal = st.number_input(f"Internal {i} (out of 40)", 0, 40, key=f"internal_{i}")
                with col4:
                    if subject_type == "Practical":
                        end = st.number_input(f"External {i} (out of 40)", 0, 40, key=f"end_{i}")
                    else:
                        end = st.number_input(f"External {i} (out of 60)", 0, 60, key=f"end_{i}")
                with col5:
                    credits = st.number_input(f"Credits {i}", 1, 6, 4, key=f"credits_{i}")
                    total = internal + end
                    if total >= 91:
                        grade = "EX"
                        gp = 10
                    elif total >= 86:
                        grade = "AA"
                        gp = 9
                    elif total >= 81:
                        grade = "AB"
                        gp = 8.5
                    elif total >= 76:
                        grade = "BB"
                        gp = 8
                    elif total >= 71:
                        grade = "BC"
                        gp = 7.5
                    elif total >= 66:
                        grade = "CC"
                        gp = 7
                    elif total >= 61:
                        grade = "CD"
                        gp = 6.5
                    elif total >= 56:
                        grade = "DD"
                        gp = 6
                    elif total >= 51:
                        grade = "DE"
                        gp = 5.5
                    elif total >= 40:
                        grade = "EE"
                        gp = 5
                    else:
                        grade = "F"
                        gp = 0
                    
                    st.write(f"**Grade: {grade}**")
                
                if subject:
                    grades_data.append({
                        "subject": subject,
                        "internal": internal,
                        "end": end,
                        "total": internal + end,
                        "grade": grade,
                        "credits": credits,
                        "gp": gp,
                        "cp": gp * credits
                    })
            
            if grades_data and st.button("Generate Grade Card", key="generate_gradecard"):
                student_data = {
                    "id": student_info['id'],
                    "name": student_info['name'],
                    "prn": student_info.get('prn', ''),
                    "roll_no": student_info['roll_no'],
                    "class": student_info['class_name'],
                    "semester": "IV",
                    "course": "B.Tech",
                    "exam_event": "",
                    "department": DEPARTMENT_NAME
                }
                
                pdf_bytes = generate_maharashtra_gradecard(student_data, grades_data)
                st.download_button(
                    "📥 Download Grade Card PDF",
                    pdf_bytes,
                    f"gradecard_{student_info['roll_no']}.pdf",
                    "application/pdf",
                    key="download_gradecard"
                )
                try:
                    word_bytes = generate_gradecard_docx(student_data, grades_data)
                    st.download_button(
                        "📥 Download Grade Card DOCX",
                        word_bytes,
                        f"gradecard_{student_info['roll_no']}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_gradecard_docx"
                    )
                except ImportError as e:
                    st.warning(str(e))
                save_gradecard(student_info['id'], pdf_bytes, student_data['semester'], student_data['course'])
                st.success("✅ Grade card generated successfully and saved for student download.")
        else:
            st.warning("Student information not found")
    
    render_page_footer()

# Main app
if 'user' not in st.session_state:
    login()
else:
    user = st.session_state.user
    st.sidebar.title(f"Welcome, {user['name']}")
    st.sidebar.button("Logout", on_click=logout)
    
    if user['role'] == 'admin':
        admin_page()
    elif user['role'] == 'faculty':
        faculty_page()
    elif user['role'] == 'student':
        student_page()